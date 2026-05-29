"""
生成《合作立项通知》Word文档
石家庄市科技政策研究会 → 石家庄金融职业学院
内部合作公文，单页
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 辅助函数 ────────────────────────────────────────────────────────────────

def set_cjk(run, east='仿宋_GB2312', west='Times New Roman',
            size=15, bold=False, color=None, italic=False):
    run.font.name = west
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    f = OxmlElement('w:rFonts')
    f.set(qn('w:ascii'),    west)
    f.set(qn('w:hAnsi'),    west)
    f.set(qn('w:eastAsia'), east)
    f.set(qn('w:cs'),       east)
    rPr.insert(0, f)
    if color:
        run.font.color.rgb = RGBColor(*color)


def para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=0, fi=0, li=0, ls=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment    = align
    fmt.space_before = Pt(sb)
    fmt.space_after  = Pt(sa)
    if fi: fmt.first_line_indent = Pt(fi)
    if li: fmt.left_indent       = Pt(li)
    if ls:
        fmt.line_spacing      = Pt(ls)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return p


def run(p, text, east='仿宋_GB2312', west='Times New Roman',
        size=15, bold=False, color=None, italic=False):
    r = p.add_run(text)
    set_cjk(r, east, west, size, bold, color, italic)
    return r


def cell_fmt(cell, bg=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             sb=3, sa=3, ls=20):
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  bg)
        tcPr.append(shd)
    p = cell.paragraphs[0]
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(sb)
    fmt.space_after  = Pt(sa)
    fmt.line_spacing = Pt(ls)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return p


def add_line(doc, which='bottom', style='single', sz='4',
             color='002F6C', sp_before=0, sp_after=4, double=False):
    p = para(doc, sb=sp_before, sa=sp_after, ls=1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{which}')
    bdr.set(qn('w:val'),   'double' if double else style)
    bdr.set(qn('w:sz'),    sz)
    bdr.set(qn('w:space'), '1')
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)


# ── 建立文档 ────────────────────────────────────────────────────────────────

doc = Document()

# 清除默认段落间距
for s in ('Normal',):
    doc.styles[s].paragraph_format.space_before = Pt(0)
    doc.styles[s].paragraph_format.space_after  = Pt(0)

# 页面（A4，上下稍压缩以保证单页）
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2.8)
sec.bottom_margin = Cm(2.6)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.8)


# ══════════════════════════════════════════════════════════════
# 版头：发文机关名称
# ══════════════════════════════════════════════════════════════

p0 = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0, ls=52)
run(p0, '石家庄市科技政策研究会',
    east='方正小标宋简体', size=26, bold=False, color=(0, 47, 108))

# 深蓝双横线
add_line(doc, which='bottom', double=True, sz='8',
         color='002F6C', sp_before=4, sp_after=6)

# 文号（左对齐）
p_no = para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=2, ls=26)
run(p_no, '石科政研〔2026〕XXX号', size=14)

# 文件种类小标（右对齐，可选）
# 不加，保持简洁


# ══════════════════════════════════════════════════════════════
# 标题
# ══════════════════════════════════════════════════════════════

p_title = para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, sb=14, sa=12, ls=36)
run(p_title,
    '关于联合承接河北省发展和改革委员会课题研究\n合作立项的通知',
    east='方正小标宋简体', size=20, bold=False)


# ══════════════════════════════════════════════════════════════
# 主送
# ══════════════════════════════════════════════════════════════

p_to = para(doc, sb=0, sa=0, ls=27)
run(p_to, '石家庄金融职业学院大数据会计学院：', size=15)


# ══════════════════════════════════════════════════════════════
# 正文引言（紧凑，2行内）
# ══════════════════════════════════════════════════════════════

p_body = para(doc, sb=0, sa=0, ls=27, fi=30)
run(p_body,
    '根据河北省发展和改革委员会2026年度城乡融合发展领域课题申报部署，'
    '我会拟以联合申报方式共同承接相关研究课题。经协商，现将合作立项事项通知如下。',
    size=15)


# ══════════════════════════════════════════════════════════════
# 一、课题基本信息
# ══════════════════════════════════════════════════════════════

p_h1 = para(doc, sb=10, sa=4, ls=27)
run(p_h1, '  一、课题基本信息', east='黑体', size=15, bold=True)

tbl = doc.add_table(rows=7, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data = [
    ('课题名称',   '河北省城乡基础设施一体化与公共服务均等化路径研究'),
    ('所属方向',   '（三）城乡融合发展与乡村全面振兴协同'),
    ('牵头单位',   '石家庄市科技政策研究会'),
    ('合作单位',   '石家庄金融职业学院大数据会计学院'),
    ('课题负责人', '王维佳（石家庄金融职业学院大数据会计学院）'),
    ('研究周期',   '2026年6月—2027年12月'),
    ('课题经费',   'XX万元（以正式合同为准）'),
]

COL0 = Cm(3.5)
for i, (label, value) in enumerate(rows_data):
    row = tbl.rows[i]
    # 标签
    c0 = row.cells[0]
    c0.width = COL0
    p0 = cell_fmt(c0, bg='EBF3FD', align=WD_ALIGN_PARAGRAPH.CENTER,
                  sb=4, sa=4, ls=21)
    run(p0, label, east='黑体', size=13, bold=True)
    # 内容
    c1 = row.cells[1]
    p1 = cell_fmt(c1, align=WD_ALIGN_PARAGRAPH.LEFT, sb=4, sa=4, ls=21)
    run(p1, value, east='仿宋_GB2312', size=13)


# ══════════════════════════════════════════════════════════════
# 二、工作要求
# ══════════════════════════════════════════════════════════════

p_h2 = para(doc, sb=10, sa=4, ls=27)
run(p_h2, '  二、工作要求', east='黑体', size=15, bold=True)

reqs = [
    ('（一）',
     '贵院请于2026年6月30日前完成合作承诺函签署，明确参与人员名单并加盖公章后报送我会。'),
    ('（二）',
     '双方按分工协议开展研究，贵院负责提供数据支撑、案例素材及子模块撰写，'
     '并参与答辩辅助等工作；具体分工以双方协商确定的协议为准。'),
    ('（三）',
     '研究成果经双方确认后联合署名，课题验收、经费使用及成果发表等事项依据'
     '后续签订的合作协议执行。'),
]

for code, content in reqs:
    p = para(doc, sb=2, sa=2, ls=27, fi=30)
    run(p, code, east='仿宋_GB2312', size=15, bold=True)
    run(p, content, size=15)


# ══════════════════════════════════════════════════════════════
# 签发区
# ══════════════════════════════════════════════════════════════

p_sp = para(doc, sb=12, sa=0, ls=12)   # 间距

p_sign1 = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=0, sa=0, ls=27)
run(p_sign1, '石家庄市科技政策研究会', size=15)

p_sign2 = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=4, sa=0, ls=27)
run(p_sign2, '2026年　　月　　日', size=15)

p_seal = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=2, sa=0, ls=20)
run(p_seal, '（加盖公章）',
    size=12, color=(160, 160, 160), italic=True)


# ══════════════════════════════════════════════════════════════
# 版记横线 + 抄送
# ══════════════════════════════════════════════════════════════

add_line(doc, which='top', style='single', sz='4',
         color='000000', sp_before=14, sp_after=4)

p_cc = para(doc, sb=0, sa=0, ls=22)
run(p_cc, '抄送：石家庄市科技政策研究会存档。',
    east='仿宋_GB2312', size=12)

p_ft = para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT, sb=2, sa=0, ls=22)
run(p_ft, '石家庄市科技政策研究会　　2026年　　月　　日印发',
    east='仿宋_GB2312', size=12)


# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════

out = '/home/user/aer-skill/合作立项通知.docx'
doc.save(out)
print(f'✓ 合作立项通知已生成：{out}')
