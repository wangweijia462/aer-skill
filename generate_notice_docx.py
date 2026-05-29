"""
生成《河北省发展和改革委员会课题立项通知》Word文档
公文格式：GB/T 9704-2012《党政机关公文格式》标准
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── 工具函数 ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_cjk(run, east='仿宋_GB2312', west='Times New Roman',
                size=16, bold=False, color=None, italic=False):
    run.font.name = west
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    # Remove existing rFonts if any
    for old in rPr.findall(qn('w:rFonts')):
        rPr.remove(old)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    west)
    rFonts.set(qn('w:hAnsi'),    west)
    rFonts.set(qn('w:eastAsia'), east)
    rFonts.set(qn('w:cs'),       east)
    rPr.insert(0, rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)


def new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
             sp_before=0, sp_after=0,
             first_indent_pt=0, left_indent_pt=0,
             line_spacing_pt=None):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.alignment      = align
    fmt.space_before   = Pt(sp_before)
    fmt.space_after    = Pt(sp_after)
    if first_indent_pt:
        fmt.first_line_indent = Pt(first_indent_pt)
    if left_indent_pt:
        fmt.left_indent = Pt(left_indent_pt)
    if line_spacing_pt:
        fmt.line_spacing      = Pt(line_spacing_pt)
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    return p


def add_run(para, text, east='仿宋_GB2312', west='Times New Roman',
            size=16, bold=False, color=None, italic=False):
    run = para.add_run(text)
    set_run_cjk(run, east, west, size, bold, color, italic)
    return run


def add_red_rule(doc, rule_type='double', thickness='6', space='1',
                 color='C40000', which='bottom', sp_before=0, sp_after=6):
    """在段落上/下添加红色横线"""
    p = new_para(doc, sp_before=sp_before, sp_after=sp_after,
                 line_spacing_pt=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr = OxmlElement(f'w:{which}')
    bdr.set(qn('w:val'),   rule_type)
    bdr.set(qn('w:sz'),    thickness)
    bdr.set(qn('w:space'), space)
    bdr.set(qn('w:color'), color)
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


def add_black_rule(doc, which='top', sp_before=0, sp_after=0):
    p = new_para(doc, sp_before=sp_before, sp_after=sp_after,
                 line_spacing_pt=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bdr  = OxmlElement(f'w:{which}')
    bdr.set(qn('w:val'),   'single')
    bdr.set(qn('w:sz'),    '6')
    bdr.set(qn('w:space'), '1')
    bdr.set(qn('w:color'), '000000')
    pBdr.append(bdr)
    pPr.append(pBdr)
    return p


# ── 建立文档 ────────────────────────────────────────────────────────────────

doc = Document()

# 清除默认段落间距
style_normal = doc.styles['Normal']
style_normal.paragraph_format.space_before = Pt(0)
style_normal.paragraph_format.space_after  = Pt(0)

# 页面设置（A4，标准公文页边距）
sec = doc.sections[0]
sec.page_width    = Cm(21.0)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(3.7)
sec.bottom_margin = Cm(3.5)
sec.left_margin   = Cm(2.8)
sec.right_margin  = Cm(2.6)


# ══════════════════════════════════════════════════════════════
# ①  版头区：发文机关标志
# ══════════════════════════════════════════════════════════════

# 顶部留白
p_top = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=24)
add_run(p_top, ' ')

# 发文机关名称（红色大字）
p_org = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                 sp_before=6, sp_after=0, line_spacing_pt=60)
add_run(p_org, '河北省发展和改革委员会',
        east='方正小标宋简体', west='Times New Roman',
        size=28, bold=False, color=(196, 0, 0))

# 文件红色双横线
add_red_rule(doc, rule_type='double', thickness='8',
             color='C40000', which='bottom', sp_before=4, sp_after=8)

# 文号
p_no = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                sp_before=6, sp_after=0, line_spacing_pt=30)
add_run(p_no, '冀发改规划〔2026〕XXX号', east='仿宋_GB2312', size=16)


# ══════════════════════════════════════════════════════════════
# ②  标题
# ══════════════════════════════════════════════════════════════

p_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                   sp_before=22, sp_after=16, line_spacing_pt=44)
add_run(p_title,
        '关于下达2026年度城乡融合发展与乡村全面振兴\n'
        '协同研究领域省级课题立项的通知',
        east='方正小标宋简体', west='Times New Roman',
        size=22, bold=False)


# ══════════════════════════════════════════════════════════════
# ③  主送机关
# ══════════════════════════════════════════════════════════════

p_recv = new_para(doc, align=WD_ALIGN_PARAGRAPH.LEFT,
                  sp_before=0, sp_after=0, line_spacing_pt=30)
add_run(p_recv, '石家庄市科技政策研究会：', east='仿宋_GB2312', size=16)


# ══════════════════════════════════════════════════════════════
# ④  正文
# ══════════════════════════════════════════════════════════════

INDENT = Pt(32)   # 首行2字符缩进（16pt × 2）
LINE   = 30       # 正文行距30pt

# 引言
p0 = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=LINE)
p0.paragraph_format.first_line_indent = INDENT
add_run(p0,
        '根据河北省发展和改革委员会《2026年度城乡融合发展领域软科学课题申报指南》'
        '（冀发改办〔2026〕XX号）有关部署，经组织专家评审，按照公平公正、择优立项的原则，'
        '对申报的省级软科学研究课题进行评审认定。现将2026年度城乡融合发展与乡村全面振兴'
        '协同研究领域立项课题情况通知如下。',
        east='仿宋_GB2312', size=16)

# ── 一、课题立项信息 ──────────────────────────────────────────
p_h1 = new_para(doc, sp_before=12, sp_after=6, line_spacing_pt=LINE)
add_run(p_h1, '  一、课题立项信息', east='黑体', size=16, bold=True)

# 信息表格
table = doc.add_table(rows=10, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

rows_data = [
    ('课题名称',   '河北省城乡基础设施一体化与公共服务均等化路径研究'),
    ('项目编号',   'HBFG2026-CXRH-XXX（以合同书为准）'),
    ('所属领域',   '城乡融合发展与乡村全面振兴协同'),
    ('研究方向',   '城乡基础设施一体化与公共服务均等化'),
    ('课题类别',   '省级软科学重点研究课题'),
    ('牵头单位',   '石家庄市科技政策研究会'),
    ('联合单位',   '石家庄金融职业学院大数据会计学院'),
    ('课题负责人', '王维佳（石家庄金融职业学院大数据会计学院）'),
    ('研究周期',   '2026年6月—2027年12月（共18个月）'),
    ('课题经费',   'XX万元（具体金额以合同书为准）'),
]

for i, (label, value) in enumerate(rows_data):
    row = table.rows[i]
    # 标签列
    c0 = row.cells[0]
    c0.width = Cm(3.8)
    set_cell_bg(c0, 'EBF3FD')
    p_c0 = c0.paragraphs[0]
    p_c0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_c0.paragraph_format.space_before = Pt(3)
    p_c0.paragraph_format.space_after  = Pt(3)
    p_c0.paragraph_format.line_spacing = Pt(22)
    p_c0.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    add_run(p_c0, label, east='黑体', size=14, bold=True)
    # 内容列
    c1 = row.cells[1]
    p_c1 = c1.paragraphs[0]
    p_c1.paragraph_format.space_before = Pt(3)
    p_c1.paragraph_format.space_after  = Pt(3)
    p_c1.paragraph_format.line_spacing = Pt(22)
    p_c1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    add_run(p_c1, value, east='仿宋_GB2312', size=14)

# ── 二、工作要求 ──────────────────────────────────────────────
p_h2 = new_para(doc, sp_before=14, sp_after=6, line_spacing_pt=LINE)
add_run(p_h2, '  二、工作要求', east='黑体', size=16, bold=True)

requirements = [
    (
        '（一）签订课题合同。',
        '请牵头单位石家庄市科技政策研究会于2026年6月30日前，持本通知及项目申报材料，'
        '与河北省发展和改革委员会规划发展处签订《课题研究合同书》，明确研究目标、阶段成果、'
        '经费使用计划及成果验收标准等事项。联合单位石家庄金融职业学院大数据会计学院须同步'
        '出具《联合参与课题研究承诺函》，并由学院加盖公章后随合同一并报送。'
    ),
    (
        '（二）落实规范管理。',
        '课题牵头单位应依据《河北省发展和改革委员会软科学研究课题管理办法》，建立课题'
        '内部管理制度，设立专项经费账户，实行专款专用，接受主管部门的财务审计和绩效评价。'
        '课题组核心研究人员须在合同中明确；研究人员如需调整，须提前向委规划发展处报批，'
        '不得擅自变更课题负责人。'
    ),
    (
        '（三）按期推进研究。',
        '课题研究分两个阶段推进：第一阶段（2026年6月—2026年12月）完成基础数据整理、'
        '实地调研及阶段性研究报告；第二阶段（2027年1月—2027年12月）完成终期研究报告、'
        '政策建议专报及学术论文不少于2篇。各阶段成果须按时报送委规划发展处，不得无故延期。'
        '如遇特殊情况需延期，须于阶段截止日30日前书面申请。'
    ),
    (
        '（四）强化成果应用。',
        '课题研究须紧密结合河北省城乡融合发展的实际需要，注重研究成果的政策可操作性'
        '与数据实证性，研究成果应直接服务于河北省"十五五"规划编制、城乡融合发展相关政策'
        '制定及省委省政府决策参考。鼓励在CSSCI、北大核心等高水平期刊公开发表研究成果，'
        '并积极参与省级高端论坛交流研讨。'
    ),
    (
        '（五）严守廉洁纪律。',
        '课题研究各单位须严格遵守科研诚信规范和学术道德准则，严禁数据造假、'
        '成果剽窃等学术不端行为。课题经费使用须严格执行国家及省有关财务管理规定，'
        '自觉接受纪检监察和审计部门的监督检查。'
    ),
]

for title, content in requirements:
    p = new_para(doc, sp_before=4, sp_after=2, line_spacing_pt=LINE)
    p.paragraph_format.first_line_indent = INDENT
    add_run(p, title, east='仿宋_GB2312', size=16, bold=True)
    add_run(p, content, east='仿宋_GB2312', size=16)

# ── 三、其他事项 ──────────────────────────────────────────────
p_h3 = new_para(doc, sp_before=14, sp_after=6, line_spacing_pt=LINE)
add_run(p_h3, '  三、其他事项', east='黑体', size=16, bold=True)

p_other = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=LINE)
p_other.paragraph_format.first_line_indent = INDENT
add_run(p_other,
        '本通知自发布之日起生效。课题研究工作中如有重大情况变化，请及时向河北省发展'
        '和改革委员会规划发展处报告。本通知未尽事宜，依据《河北省发展和改革委员会软科学'
        '研究课题管理办法》有关规定执行。业务联系人：XXX，联系电话：0311-XXXXXXXX，'
        '电子邮箱：XXXXX@hbdrc.gov.cn。',
        east='仿宋_GB2312', size=16)

# 附件说明
p_att_label = new_para(doc, sp_before=12, sp_after=2, line_spacing_pt=LINE)
add_run(p_att_label, '  附件：', east='仿宋_GB2312', size=16)

p_att1 = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=LINE,
                  left_indent_pt=48)
add_run(p_att1,
        '1．2026年度城乡融合发展与乡村全面振兴协同研究领域省级课题立项清单',
        east='仿宋_GB2312', size=16)
p_att2 = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=LINE,
                  left_indent_pt=48)
add_run(p_att2,
        '2．课题研究合同书（样本）',
        east='仿宋_GB2312', size=16)


# ══════════════════════════════════════════════════════════════
# ⑤  版记：发文机关署名 + 日期
# ══════════════════════════════════════════════════════════════

# 机关署名右对齐
p_sign1 = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT,
                   sp_before=28, sp_after=0, line_spacing_pt=LINE)
add_run(p_sign1, '河北省发展和改革委员会', east='仿宋_GB2312', size=16)

# 成文日期右对齐（手动填写）
p_sign2 = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT,
                   sp_before=4, sp_after=0, line_spacing_pt=LINE)
add_run(p_sign2, '2026年　　月　　日', east='仿宋_GB2312', size=16)

# 盖章提示（小字斜体）
p_seal = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT,
                  sp_before=4, sp_after=0, line_spacing_pt=22)
add_run(p_seal, '（此处加盖河北省发展和改革委员会印章）',
        east='仿宋_GB2312', size=12,
        color=(160, 160, 160), italic=True)


# ══════════════════════════════════════════════════════════════
# ⑥  版记区：黑色单横线 + 抄送 + 印发机关
# ══════════════════════════════════════════════════════════════

add_black_rule(doc, which='top', sp_before=20, sp_after=4)

p_cc = new_para(doc, sp_before=0, sp_after=0, line_spacing_pt=24)
add_run(p_cc,
        '抄送：河北省发展和改革委员会规划发展处，石家庄金融职业学院大数据会计学院。',
        east='仿宋_GB2312', size=14)

p_footer = new_para(doc, align=WD_ALIGN_PARAGRAPH.RIGHT,
                    sp_before=2, sp_after=0, line_spacing_pt=24)
add_run(p_footer,
        '河北省发展和改革委员会办公室　　2026年　　月　　日印发',
        east='仿宋_GB2312', size=14)

# 页底分隔线
add_black_rule(doc, which='bottom', sp_before=4, sp_after=0)


# ══════════════════════════════════════════════════════════════
# ⑦  附件（另起一页）
# ══════════════════════════════════════════════════════════════

doc.add_page_break()

p_att_h = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                   sp_before=20, sp_after=16, line_spacing_pt=38)
add_run(p_att_h, '附件1',
        east='方正小标宋简体', size=18, bold=False)

p_att_title = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                       sp_before=0, sp_after=20, line_spacing_pt=38)
add_run(p_att_title,
        '2026年度城乡融合发展与乡村全面振兴协同研究\n领域省级课题立项清单',
        east='方正小标宋简体', size=18, bold=False)

# 附件表格
att_table = doc.add_table(rows=4, cols=5)
att_table.style = 'Table Grid'
att_table.alignment = WD_TABLE_ALIGNMENT.CENTER

att_headers = ['序号', '课题名称', '牵头单位', '课题负责人', '研究周期']
att_widths   = [Cm(1.2), Cm(6.5), Cm(4.0), Cm(2.5), Cm(2.5)]

# 表头
hdr_row = att_table.rows[0]
for j, (h, w) in enumerate(zip(att_headers, att_widths)):
    cell = hdr_row.cells[j]
    cell.width = w
    set_cell_bg(cell, '002F6C')
    p_h = cell.paragraphs[0]
    p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_h.paragraph_format.space_before = Pt(4)
    p_h.paragraph_format.space_after  = Pt(4)
    p_h.paragraph_format.line_spacing = Pt(20)
    p_h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    add_run(p_h, h, east='黑体', size=13, bold=True, color=(255, 255, 255))

# 数据行1：本课题
att_row1 = [
    '1',
    '河北省城乡基础设施一体化与公共服务均等化路径研究',
    '石家庄市科技政策研究会',
    '王维佳',
    '2026.6—2027.12',
]
for j, val in enumerate(att_row1):
    cell = att_table.rows[1].cells[j]
    p_v = cell.paragraphs[0]
    p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER if j != 1 else WD_ALIGN_PARAGRAPH.LEFT
    p_v.paragraph_format.space_before = Pt(3)
    p_v.paragraph_format.space_after  = Pt(3)
    p_v.paragraph_format.line_spacing = Pt(20)
    p_v.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    add_run(p_v, val, east='仿宋_GB2312', size=12)

# 其余行留空（供其他课题填写）
for r_idx in [2, 3]:
    for j in range(5):
        cell = att_table.rows[r_idx].cells[j]
        p_v = cell.paragraphs[0]
        p_v.paragraph_format.space_before = Pt(3)
        p_v.paragraph_format.space_after  = Pt(3)
        p_v.paragraph_format.line_spacing = Pt(20)
        p_v.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        add_run(p_v, '', east='仿宋_GB2312', size=12)

p_note = new_para(doc, sp_before=10, sp_after=0, line_spacing_pt=24)
add_run(p_note,
        '注：本清单仅列本次下达课题，如有其他同批次课题另行通知。',
        east='仿宋_GB2312', size=12,
        color=(100, 100, 100), italic=True)


# ══════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════

output_path = '/home/user/aer-skill/河北省发改委课题立项通知.docx'
doc.save(output_path)
print(f'✓ 立项通知已生成：{output_path}')
